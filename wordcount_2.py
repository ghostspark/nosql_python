import docx

file=docx.Document(input("请输入文件以及其路径："))
i=len(file.paragraphs)
with open("tar.txt","w",encoding='utf-8') as f:
    for para in range(i):
        f.write(file.paragraphs[para].text)
    f.close()

    
